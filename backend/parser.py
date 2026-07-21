import io
import string
import pdfplumber
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from a PDF file in-memory using pdfplumber."""
    text_content = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    return "\n".join(text_content)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from a DOCX file in-memory using python-docx."""
    doc = docx.Document(io.BytesIO(file_bytes))
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text_content.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text_content.append(" | ".join(row_text))
                
    return "\n".join(text_content)

def validate_extracted_text(text: str, is_pdf: bool = False) -> dict:
    """
    Validates the quality of the extracted text.
    Returns a dict with key 'status' as 'valid', 'scanned_pdf', 'too_short', or 'garbled', 
    and a descriptive 'message'.
    """
    cleaned_text = text.strip()
    
    # Check if empty or extremely short
    if len(cleaned_text) < 50:
        if is_pdf:
            return {
                "status": "scanned_pdf",
                "message": "The PDF file appears to be scanned or image-only. Fallback to image-based analysis is required."
            }
        else:
            return {
                "status": "too_short",
                "message": "The extracted text is too short. Please provide a more detailed resume."
            }
            
    # Check for garbled/unreadable text (proportion of printable ASCII characters)
    printable_chars = set(string.printable)
    printable_count = sum(1 for char in cleaned_text if char in printable_chars)
    printable_ratio = printable_count / len(cleaned_text)
    
    if printable_ratio < 0.85:
        return {
            "status": "garbled",
            "message": "The document's text encoding is unreadable. Please upload a clear PDF, DOCX, or image file."
        }
        
    return {
        "status": "valid",
        "message": "Text validation successful."
    }

def convert_pdf_to_images(pdf_bytes: bytes, max_pages: int = 3) -> list[bytes]:
    """
    Converts PDF pages into PNG image bytes in-memory using pypdfium2.
    Limits conversion to max_pages to prevent heavy payloads.
    """
    import pypdfium2 as pdfium
    
    images = []
    try:
        pdf = pdfium.PdfDocument(pdf_bytes)
        num_pages = len(pdf)
        for i in range(min(num_pages, max_pages)):
            page = pdf[i]
            # Render at 150 DPI for reasonable readability and file size
            bitmap = page.render(scale=150 / 72.0)
            pil_image = bitmap.to_pil()
            
            buffer = io.BytesIO()
            pil_image.save(buffer, format="PNG")
            images.append(buffer.getvalue())
            
            page.close()
        pdf.close()
    except Exception as e:
        print(f"Error converting PDF to images: {str(e)}")
    return images

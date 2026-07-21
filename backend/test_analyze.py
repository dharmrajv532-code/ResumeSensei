import docx
import httpx
import json
import os

def create_mock_docx(filename: str):
    """Creates a mock resume in DOCX format to test parsing and analysis."""
    doc = docx.Document()
    
    # Title
    doc.add_heading('John Doe', 0)
    
    # Contact Info
    doc.add_paragraph('john.doe@example.com | (555) 019-2834 | San Francisco, CA')
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        'Experienced Software Engineer with a demonstrated history of working in the software development industry. '
        'Skilled in Python, JavaScript, and cloud technologies. Strong engineering professional with a Bachelor\'s degree.'
    )
    
    # Experience
    doc.add_heading('Work Experience', level=1)
    
    # Job 1
    doc.add_heading('Senior Software Engineer - Tech Solutions Inc.', level=2)
    doc.add_paragraph('January 2022 - Present')
    doc.add_paragraph('• Responsible for developing new features for the main application.')
    doc.add_paragraph('• Worked with the team to fix bugs and improve performance.')
    doc.add_paragraph('• Helped migrate legacy systems to the cloud, making things faster.')
    
    # Job 2
    doc.add_heading('Software Engineer - CodeCraft Corp', level=2)
    doc.add_paragraph('June 2019 - December 2021')
    doc.add_paragraph('• Wrote code in Python and Javascript.')
    doc.add_paragraph('• Participated in daily standups and agile sprint planning sessions.')
    doc.add_paragraph('• Wrote unit tests for new code to improve reliability.')

    # Skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Programming Languages: Python, JavaScript, SQL, HTML, CSS')
    doc.add_paragraph('Frameworks & Tools: FastAPI, React, Node.js, Docker, AWS')
    doc.add_paragraph('Soft Skills: Teamwork, communication, problem-solving')

    # Save
    doc.save(filename)
    print(f"Mock resume created at: {filename}")

def test_analyze_endpoint(filename: str):
    """Submits the mock resume to the running FastAPI analyze endpoint."""
    url = "http://127.0.0.1:8000/analyze"
    print(f"Sending request to: {url}...")
    
    try:
        with open(filename, "rb") as f:
            files = {"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            # 60s timeout to allow Groq LLM processing
            response = httpx.post(url, files=files, timeout=60.0)
            
        print(f"Status Code: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print("\n--- ANALYSIS RESPONSE SUCCEEDED ---\n")
            print(f"File Name: {result.get('file_name')}")
            print(f"File Type: {result.get('file_type')}")
            print(f"Status: {result.get('status')}")
            print(f"Message: {result.get('message')}")
            
            analysis = result.get('analysis')
            if analysis:
                print(f"\nOverall Score: {analysis.get('overall_score')}/100")
                print("\nCategory Scores:")
                for cat, score in analysis.get('category_scores', {}).items():
                    print(f"  - {cat}: {score}")
                
                print("\nExtracted Skills:")
                print(f"  - Technical: {analysis.get('extracted_skills', {}).get('technical', [])}")
                print(f"  - Soft: {analysis.get('extracted_skills', {}).get('soft', [])}")
                
                print(f"\nMissing Sections: {analysis.get('missing_sections', [])}")
                
                print("\nIssues Identified:")
                for issue in analysis.get('issues', [])[:3]:
                    print(f"  - [{issue.get('severity')}] {issue.get('section')}: {issue.get('problem')}")
                
                print("\nActionable Recommendations:")
                for rec in analysis.get('recommendations', [])[:3]:
                    print(f"  - {rec}")
                    
                print("\nBullet point rewrites sample:")
                for rewrite in analysis.get('bullet_rewrites', [])[:2]:
                    print(f"  Original: {rewrite.get('original')}")
                    print(f"  Improved: {rewrite.get('improved')}")
                    print(f"  Reason  : {rewrite.get('reason')}\n")
            else:
                print("No analysis payload returned.")
        else:
            print(f"Error details: {response.text}")
            
    except Exception as e:
        print(f"Failed to connect or test: {str(e)}")

if __name__ == "__main__":
    resume_file = "john_doe_resume.docx"
    create_mock_docx(resume_file)
    test_analyze_endpoint(resume_file)
    
    # Cleanup
    if os.path.exists(resume_file):
        os.remove(resume_file)
        print("Cleaned up temporary mock file.")
